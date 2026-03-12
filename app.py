import streamlit as st
import requests
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

# --- CONFIGURAÇÕES ---
minha_chave = "AIzaSyCIPi18CGzOPtSiJJHHRGXtYI2nGJYdcc0"

st.set_page_config(page_title="Gerador UFAM - Vangles", layout="wide")

# --- SUA LÓGICA DE IA ---
def chamar_minha_ia(prompt_comando):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={minha_chave}"
    headers = {'Content-Type': 'application/json'}
    prompt_final = (f"{prompt_comando} REGRAS: Inicie o texto diretamente. Use linguagem acadêmica e impessoal. "
                    f"IMPORTANTE: Se usar citações, ao final do texto, escreva a tag [REFS] e coloque as referências completas em ABNT.")
    corpo = {"contents": [{"parts": [{"text": prompt_final}]}]}
    try:
        response = requests.post(url, headers=headers, data=json.dumps(corpo))
        return response.json()['candidates'][0]['content']['parts'][0]['text']
    except: return "Erro na conexão com a IA."

# --- SUA FORMATAÇÃO WORD (PRESERVADA) ---
def add_estilo(doc, texto, negrito=False, tam=12, alinhar=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo_esq=0, recuo_primeira_linha=0, eh_simples=False):
    paragrafos = texto.split('\n')
    for trecho in paragrafos:
        trecho = trecho.strip()
        if not trecho: continue
        p = doc.add_paragraph()
        p.alignment = alinhar
        if recuo_esq > 0: p.paragraph_format.left_indent = Cm(recuo_esq)
        if recuo_primeira_linha > 0 and not eh_simples: p.paragraph_format.first_line_indent = Cm(recuo_primeira_linha)
        p.paragraph_format.line_spacing = 1.0 if eh_simples else 1.5
        run = p.add_run(trecho)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tam)
        run.bold = negrito

# --- INTERFACE ---
st.title("🎓 Gerador de Relatório UFAM")
st.markdown("---")

if 'dados' not in st.session_state:
    st.session_state.dados = {"metodologia": "", "introducao": "", "titulo": "Relatório de Estágio"}

col1, col2 = st.columns(2)
with col1:
    st.subheader("Dados da Escola")
    escola = st.text_input("Nome da Escola", "Escola Estadual Indígena Professor Gildo Sampaio Megatanücü")
    local = st.text_input("Localização", "na zona rural de Benjamin Constant-AM")
    turma = st.text_input("Turma", "101")
    tema = st.text_input("Tema ODS", "ODS 15 Vida Terrestre")

with col2:
    st.subheader("Relato da Intervenção")
    relato_raw = st.text_area("O que aconteceu na aula? (Para a IA ajustar)", height=150)

if st.button("🚀 Gerar Metodologia Completa"):
    with st.spinner("Formatando Etapas I a VI..."):
        res_ia = chamar_minha_ia(relato_raw)
        
        # SUA F-STRING INTEGRAL COM TODAS AS ETAPAS
        texto_completo = f"""A carga horária total do estágio é de 135 horas, distribuídas em seis (6) etapas:

**I – Aulas Teóricas (45 horas):**
Abordando as normativas que regem o estágio no curso de Licenciatura em Ciências: Biologia e Química...

**II – Observação Escolar (30 horas):**
Realizada na "{escola}", localizada {local}...

**III – Planejamento Pedagógico (20 horas):**
Elaboração de atividades e planos de aula...

**IV – Regência e Intervenção (20 horas):**
Atuação na turma {turma} com o tema {tema}. {res_ia}

**V – Elaboração do Relatório Final (15 horas):**
Sistematização dos resultados e reflexões...

**VI – Apresentação dos resultados (5 horas):**
Preparação de Banner para compartilhamento das experiências."""
        
        st.session_state.dados['metodologia'] = texto_completo
        st.success("Tudo pronto!")
        st.markdown(texto_completo)

st.markdown("---")
if st.session_state.dados['metodologia']:
    if st.button("📥 Baixar Relatório em Word"):
        doc = Document()
        add_estilo(doc, "PROCEDIMENTOS METODOLÓGICOS", True, alinhar=WD_ALIGN_PARAGRAPH.CENTER)
        add_estilo(doc, st.session_state.dados['metodologia'])
        
        buffer = BytesIO()
        doc.save(buffer)
        st.download_button("Clique para Baixar", buffer.getvalue(), "Relatorio_UFAM.docx")
