import requests
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- CONFIGURAÇÕES ---
minha_chave = os.environ.get("GEMINI_API_KEY")
ARQUIVO_MEMORIA = "projeto_estagio_vangles.json"

# --- MOTOR DE IA ---
def encontrar_melhor_modelo():
    url_lista = f"https://generativelanguage.googleapis.com/v1beta/models?key={minha_chave}"
    try:
        res = requests.get(url_lista).json()
        modelos = [m['name'] for m in res.get('models', []) if 'generateContent' in m.get('supportedGenerationMethods', [])]
        for m in modelos:
            if '1.5-flash' in m: return m
        return modelos[0] if modelos else None
    except: return None

def chamar_minha_ia(prompt_comando, modelo_nome):
    url = f"https://generativelanguage.googleapis.com/v1beta/{modelo_nome}:generateContent?key={minha_chave}"
    headers = {'Content-Type': 'application/json'}
    prompt_final = (f"{prompt_comando} REGRAS: Inicie o texto diretamente. Use linguagem acadêmica e impessoal. "
                    f"IMPORTANTE: Se usar citações, ao final do texto, escreva a tag [REFS] e coloque as referências completas em ABNT.")
    corpo = {"contents": [{"parts": [{"text": prompt_final}]}]}
    try:
        response = requests.post(url, headers=headers, data=json.dumps(corpo))
        return response.json()['candidates'][0]['content']['parts'][0]['text']
    except: return "Erro na conexão."

# --- FUNÇÕES DE DADOS ---
def carregar_dados():
    if os.path.exists(ARQUIVO_MEMORIA):
        with open(ARQUIVO_MEMORIA, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {"titulo_trabalho":"","autor":"","orientador":"","cidade":"","ano":"","resumo":"","introducao":"","referencial":"","metodologia":"","resultados":"","conclusao":"","referencias":"","titulo_secao_2":"2. REFERENCIAL TEÓRICO"}

def salvar_dados(dados):
    with open(ARQUIVO_MEMORIA, 'w', encoding='utf-8') as f:
        json.dump(dados, f, indent=4, ensure_ascii=False)

def processar_e_extrair_refs(resposta_ia, dados):
    texto_limpo = resposta_ia
    refs_detectadas = []
    meses = ["jan.", "fev.", "mar.", "abr.", "mai.", "jun.", "jul.", "ago.", "set.", "out.", "nov.", "dez."]
    hoje = datetime.now()
    data_formatada = f"{hoje.day:02d} {meses[hoje.month-1]} {hoje.year}"
    
    if "[REFS]" in resposta_ia:
        partes = resposta_ia.split("[REFS]")
        texto_limpo = partes[0].strip()
        cruas = partes[1].strip().split('\n')
        for r in cruas:
            if r.strip():
                r_limpa = r.replace('*', '').strip()
                if "Acesso em:" in r_limpa:
                    base = r_limpa.split("Acesso em:")[0]
                    r_limpa = f"{base}Acesso em: {data_formatada}."
                refs_detectadas.append(r_limpa)
    return texto_limpo, refs_detectadas

# --- FORMATAÇÃO WORD ---
def add_page_number(paragraph):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def add_estilo(doc, texto, negrito=False, tam=12, alinhar=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo_esq=0, recuo_primeira_linha=0, eh_simples=False):
    paragrafos = texto.split('\n')
    for trecho in paragrafos:
        trecho = trecho.strip()
        if not trecho: continue
        
        if trecho in ["RESUMO", "SUMÁRIO"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(trecho); run.bold = True
            run.font.size, run.font.name = Pt(12), 'Times New Roman'
            continue

        eh_legenda = trecho.startswith("Figura 01:") or trecho.startswith("Fonte:") or trecho == "(LOCAL DA FOTO)"

        import re
        eh_etapa_romana = bool(re.match(r'^\*\*(I{1,3}V?|VI?I?I?|IV|IX|X)\s*[\u2013\-]', trecho))
        trecho_limpo_bold = trecho.replace('**', '').strip()

        partes = trecho.split()
        primeira = partes[0] if partes else ""
        eh_num = primeira and primeira[0].isdigit() and "." in primeira
        nivel = None
        if eh_num and not eh_legenda:
            pontos = primeira.count('.')
            if pontos == 1:
                nivel = 1 if primeira.endswith('.') else 2
            elif pontos >= 2:
                nivel = 3

        p = doc.add_heading("", level=nivel) if nivel else doc.add_paragraph()

        if eh_legenda:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.first_line_indent = 0
        elif nivel:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = 0
            p.paragraph_format.space_before = Pt(24)
        elif eh_etapa_romana:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = 0
            p.paragraph_format.left_indent = 0
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(12)
        else:
            p.alignment = alinhar
            if recuo_esq > 0: p.paragraph_format.left_indent = Cm(recuo_esq)
            if recuo_primeira_linha > 0 and not eh_simples: p.paragraph_format.first_line_indent = Cm(recuo_primeira_linha)
            p.paragraph_format.line_spacing = 1.0 if eh_simples else 1.5

        texto_renderizar = trecho_limpo_bold if eh_etapa_romana else trecho
        run = p.add_run(texto_renderizar)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0,0,0)

        if eh_legenda:
            run.font.size = Pt(10)
            run.bold = True if trecho.startswith("Figura") else False
        elif eh_etapa_romana:
            run.font.size = Pt(tam)
            run.bold = True
        else:
            run.font.size = Pt(tam)
            if nivel == 1: run.bold = True
            elif nivel == 2: run.bold, run.italic = False, False
            else: run.bold = negrito

def gerar_word_completo(dados):
    try:
        doc = Document()
        for s in doc.sections: s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(3), Cm(2), Cm(3), Cm(2)
        add_estilo(doc, "UNIVERSIDADE FEDERAL DO AMAZONAS - UFAM\nINSTITUTO NATUREZA E CULTURA – INC\nCURSO DE LICENCIATURA EM CIÊNCIAS: BIOLOGIA E QUÍMICA", True, alinhar=1, eh_simples=True)
        for _ in range(12): doc.add_paragraph()
        add_estilo(doc, dados['titulo_trabalho'].upper(), True, alinhar=1)
        for _ in range(16): doc.add_paragraph()
        add_estilo(doc, f"{dados['cidade'].upper()}\n{dados['ano']}", False, alinhar=1)
        doc.add_page_break()
        add_estilo(doc, dados['autor'].upper(), True, alinhar=1)
        for _ in range(10): doc.add_paragraph()
        add_estilo(doc, dados['titulo_trabalho'].upper(), True, alinhar=1)
        texto_apre = f"Relatório apresentado ao Instituto de Natureza e Cultura – INC/UFAM, como requisito parcial para obtenção de nota na disciplina de Estágio."
        add_estilo(doc, texto_apre, tam=10, recuo_esq=8, eh_simples=True)
        for _ in range(3): doc.add_paragraph()
        add_estilo(doc, f"Orientador (a): {dados['orientador']}", False, alinhar=1)
        for _ in range(8): doc.add_paragraph()
        add_estilo(doc, f"{dados['cidade'].upper()}\n{dados['ano']}", False, alinhar=1)
        doc.add_page_break()
        if dados.get('resumo'):
            add_estilo(doc, "RESUMO"); add_estilo(doc, dados['resumo'], recuo_primeira_linha=1.25); doc.add_page_break()
        add_estilo(doc, "SUMÁRIO")
        p_sum = doc.add_paragraph(); p_sum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_sum = p_sum.add_run()
        fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin'); run_sum._r.append(fld)
        ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'), 'preserve'); ins.text = 'TOC \\o "1-3" \\h \\z \\u'; run_sum._r.append(ins)
        fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end'); run_sum._r.append(fld2)
        doc.add_page_break()
        new_sec = doc.add_section(); new_sec.footer.is_linked_to_previous = False
        tit2 = dados.get('titulo_secao_2', "2. REFERENCIAL TEÓRICO")
        if not tit2.startswith("2. "): tit2 = f"2. {tit2}"
        secoes = [("1. INTRODUÇÃO", dados['introducao'], False), (tit2, dados['referencial'], False), ("3. PROCEDIMENTOS METODOLÓGICOS", dados['metodologia'], False), ("4. RESULTADOS E DISCUSSÕES", dados['resultados'], False), ("5. CONSIDERAÇÕES FINAIS", dados['conclusao'], False), ("6. REFERÊNCIAS BIBLIOGRÁFICAS", dados['referencias'], True)]
        for t, c, s in secoes:
            if c:
                add_estilo(doc, t, True); add_estilo(doc, c, recuo_primeira_linha=1.25, eh_simples=s)
                add_page_number(new_sec.footer.paragraphs[0]); doc.add_page_break()
        nome = f"Relatorio_{dados['autor'].replace(' ', '_') if dados['autor'] else 'Estagio'}.docx"
        doc.save(nome); print(f"✓ Word Gerado: {nome}")
    except Exception as e: print(f"Erro: {e}")

# --- CONFIGURAÇÃO INICIAL DA CAPA E FOLHA DE ROSTO ---
def configurar_capa(dados):
    campos_capa = [
        ("titulo_trabalho", "CAPA — Título do Relatório"),
        ("cidade",          "CAPA — Cidade"),
        ("ano",             "CAPA — Ano"),
        ("autor",           "FOLHA DE ROSTO — Nome completo do Autor"),
        ("orientador",      "FOLHA DE ROSTO — Nome completo do Orientador(a)"),
    ]
    print("\n" + "="*65)
    print("  CONFIGURAÇÃO INICIAL — CAPA E FOLHA DE ROSTO")
    print("="*65)
    alterou = False
    for chave, label in campos_capa:
        atual = dados.get(chave, "")
        if atual:
            print(f"\n{label}")
            print(f"  Atual: {atual}")
            resp = input("  Deseja alterar? (s/n): ").lower()
            if resp != 's':
                continue
        while True:
            valor = input(f"\n{label}:\n> ").strip()
            if not valor:
                print("  Campo obrigatório! Digite um valor.")
                continue
            conf = input(f'  Confirma: "{valor}"? (s/n): ').lower()
            if conf == 's':
                dados[chave] = valor
                alterou = True
                break
    if alterou:
        salvar_dados(dados)
        print("\n✓ Capa e Folha de Rosto configuradas com sucesso!")
    print("="*65)
    return dados

# --- MENU PRINCIPAL ---
dados = carregar_dados()
modelo_ativo = encontrar_melhor_modelo()

# Verifica se os dados da capa estão preenchidos
campos_obrigatorios = ['titulo_trabalho', 'autor', 'orientador', 'cidade', 'ano']
if not all(dados.get(c) for c in campos_obrigatorios):
    print("\n⚠ Dados da capa/folha de rosto incompletos. Vamos configurar agora.")
    dados = configurar_capa(dados)

while True:
    tags = {k: "OK" if v else "  " for k, v in dados.items()}
    capa_ok = "OK" if all(dados.get(c) for c in ["titulo_trabalho","autor","orientador","cidade","ano"]) else "  "
    print(f"\n=== MENU UFAM ===")
    print(f"C. Capa / Folha de Rosto  [{capa_ok}]")
    print(f"0. Resumo                 [{tags.get('resumo','  ')}]    1. Introdução               [{tags.get('introducao','  ')}]")
    print(f"2. Seção 2 (Referencial)   [{tags.get('referencial','  ')}]    3. Procedimentos Metodol.  [{tags.get('metodologia','  ')}]")
    print(f"4. Resultados e Discussões [{tags.get('resultados','  ')}]    5. Considerações Finais    [{tags.get('conclusao','  ')}]")
    print(f"6. Referências Bibliogr.   [{tags.get('referencias','  ')}]")
    print("-" * 65)
    print("7. GERAR WORD | 8. SAIR | 9. EXCLUIR TUDO (RECOMEÇAR)")
    
    op = input("\nEscolha: ")
    mapa = {"0":"resumo", "1":"introducao", "4":"resultados", "5":"conclusao", "6":"referencias"}
    
    if op in mapa or op in ["2", "3"]:
        chave = mapa.get(op, "referencial" if op == "2" else "metodologia")
        if dados[chave]:
            print(f"\nA seção {chave.upper()} já possui conteúdo.")
            sub = input("[R] Refazer do zero | [C] Corrigir/Adicionar | [V] Voltar: ").lower()
            if sub == 'v': continue
            elif sub == 'r': dados[chave] = ""

        if op == "2":
            while True:
                t2 = input("\nDigite o Título Principal da Seção 2:\n> ")
                if input(f"Confirma Título: {t2}? (s/n): ").lower() == 's': dados['titulo_secao_2'] = t2; break
            acumulado = ""
            for i in ["2.1", "2.2", "2.3"]:
                while True:
                    sub_t = input(f"\nNome do Subtítulo {i}:\n> ")
                    if input(f"Confirma {i} {sub_t}? (s/n): ").lower() == 's': break
                while True:
                    ped = input(f"O que a IA escreve para '{sub_t}'? Descreva o conteúdo desejado:\n> ")
                    if input("Confirma pedido? (s/n): ").lower() == 's':
                        print(f"\nProcessando {i} - {sub_t}... Aguarde.")

                        # ✅ CORREÇÃO PRINCIPAL: prompt muito mais detalhado e explícito
                        prompt_referencial = (
                            f"Escreva um texto acadêmico completo com no mínimo 4 parágrafos bem desenvolvidos "
                            f"para a subseção '{i} {sub_t}', que faz parte do Referencial Teórico sobre o tema geral "
                            f"'{dados['titulo_secao_2']}'. "
                            f"Instruções adicionais do autor: {ped}. "
                            f"IMPORTANTE: NÃO escreva o título da subseção no início. "
                            f"Escreva APENAS o texto corrido em linguagem acadêmica, com desenvolvimento teórico, "
                            f"citações de autores e embasamento científico. "
                            f"O texto deve ser completo e substancial, não apenas uma introdução."
                        )

                        res = chamar_minha_ia(prompt_referencial, modelo_ativo)

                        # ✅ CORREÇÃO SECUNDÁRIA: debug para ver o que a IA retornou
                        print(f"\n--- TEXTO GERADO PELA IA ---\n{res}\n----------------------------")

                        texto_l, refs_l = processar_e_extrair_refs(res, dados)

                        # ✅ CORREÇÃO TERCIÁRIA: remove título só se a primeira linha for APENAS o título
                        # (evita remover conteúdo que começa com o número da seção)
                        linhas_texto = texto_l.split('\n')
                        primeira_linha = linhas_texto[0].strip()
                        # Remove a primeira linha somente se ela for exatamente o subtítulo repetido
                        if primeira_linha.lower().startswith(i) and len(primeira_linha) < 100:
                            texto_l = '\n'.join(linhas_texto[1:]).strip()

                        decisao_ref = input("\n[S] Salvar | [P] Gerar nova versão | [N] Descartar e voltar: ").lower()
                        if decisao_ref == 's':
                            acumulado += f"\n{i} {sub_t.upper()}\n{texto_l}\n"
                            refs_atuais = dados['referencias'].split('\n')
                            for r in refs_l:
                                if r not in refs_atuais: refs_atuais.append(r)
                            dados['referencias'] = "\n".join(filter(None, sorted(refs_atuais)))
                            salvar_dados(dados)
                            print(f"\n✓ Subseção {i} salva com sucesso!")
                            break
                        elif decisao_ref == 'p':
                            print("\nGerando nova versão...\n")
                            continue
                        else:
                            print("\nDescartado. Tente novamente.")
                            break

            dados['referencial'] = acumulado
        
        elif op == "3":
            print("\n--- PREENCHIMENTO: PROCEDIMENTOS METODOLÓGICOS ---")
            def coletar(prompt_msg, padrao):
                while True:
                    resp = input(f"{prompt_msg} (Padrão: {padrao})\n> ") or padrao
                    conf = input(f"Confirmar '{resp}'? [S] Salvar | [N] Editar | [C] Cancelar: ").lower()
                    if conf == 's': return resp
                    if conf == 'c': return None
            
            estagio_nome = coletar("Nome do Estágio", "Estágio Supervisionado no Ensino de Biologia")
            if not estagio_nome: continue
            disciplina = coletar("Disciplina", "Biologia")
            if not disciplina: continue
            escola = coletar("Nome da Escola", "Escola Estadual Indígena Professor Gildo Sampaio Megatanücü")
            if not escola: continue
            local_escola = coletar("Localização da Escola", "na comunidade Aldeia Indígena Filadelfia, na zona rural do município de Benjamin Constant-AM")
            if not local_escola: continue
            legenda_fig = coletar("Legenda da Figura 01", "Parte da frente da Escola")
            if not legenda_fig: continue
            fonte_fig = coletar("Fonte da Figura", "Autor, 2025")
            if not fonte_fig: continue
            turma_int = coletar("Turma da Intervenção", "101")
            if not turma_int: continue
            tema_int = coletar("Tema da Intervenção", "ODS 15 Vida Terrestre")
            if not tema_int: continue

            print("\nDescreva a dinâmica/jogo da intervenção (A IA ajustará apenas este trecho):")
            rel_i = input("> ")
            
            prompt_ia = f"Com base no relato: {rel_i}. Escreva um parágrafo acadêmico curto para a seção de regência, sem títulos."
            print("Ajustando relato com IA...")
            texto_regencia = chamar_minha_ia(prompt_ia, modelo_ativo)

            texto_final = f"""O "{estagio_nome}" compreende atividades de observação no contexto escolar, bem como o acompanhamento e a prática da docência na disciplina de "{disciplina}" do 1º ao 3º ano do Ensino Médio. Inclui ainda o desenvolvimento do Planejamento de Ensino, Avaliação, Regência e Intervenção.

Na Figura 01 podemos observar a "{escola}".
Figura 01: {legenda_fig}
(LOCAL DA FOTO)
Fonte: {fonte_fig}.

A carga horária total do estágio é de 135 horas, distribuídas em seis (6) etapas, conforme descrito a seguir:

**I – Aulas Teóricas (45 horas):**
Inicialmente, foram ministradas aulas introdutórias abordando as normativas que regem o estágio no curso de Licenciatura em Ciências: Biologia e Química. Ao longo dos encontros teóricos, discutiram-se temas essenciais à formação docente, tais como: a relevância do estágio na formação de professores, a articulação entre teoria e prática, e os fundamentos da formação inicial de educadores. Durante essa fase, realizaram-se o preenchimento dos documentos principais do estágio, além de métodos pedagógicos para utilizar na prática, com o propósito de fixar os conteúdos explorados em sala.

**II – Observação Escolar (30 horas):**
Esta etapa teve como objetivo a inserção dos licenciandos na realidade escolar. Consistiu na análise do espaço físico, da estrutura organizacional, do funcionamento da instituição, da proposta pedagógica e das relações interpessoais estabelecidas no ambiente educacional. A escola escolhida foi a "{escola}", localizada {local_escola}. A observação foi realizada nos seguintes espaços: setores administrativos, biblioteca, laboratório de Ciências, refeitório, áreas externas e salas de aula, com dedicação de 5 horas para cada setor.

**III – Planejamento Pedagógico (20 horas):**
A etapa de planejamento consistiu na elaboração de atividades que seriam aplicadas durante a regência. Após a definição das turmas em que os estagiários atuariam, o professor responsável pela disciplina de "{disciplina}" indicou os conteúdos a serem trabalhados. Com base nisso, foram construídos o plano de estágio, os planos de aula e o projeto de intervenção pedagógica.

**IV – Regência e Intervenção (20 horas):**
Neste momento, realizaram-se 15 horas de regência efetiva em sala e 5 horas destinadas à execução do plano de intervenção. A atuação docente ocorreu com todas as turmas do 101 ao 302 do ensino Médio. As aulas foram conduzidas por meio de metodologias diversificadas, incluindo o uso do quadro branco, vídeos didáticos e interação com os estudantes, caracterizando-se como aulas expositivo-dialogadas. A intervenção foi aplicada na turma do "{turma_int}, com o tema {tema_int}". {texto_regencia}

**V – Elaboração do Relatório Final (15 horas):**
Essa fase consistiu na construção do relatório descritivo do estágio. Foram realizadas pesquisas bibliográficas sobre a temática "{dados['titulo_trabalho'] if dados['titulo_trabalho'] else 'formação docente'}", as quais fundamentaram teoricamente a escrita do relatório. Também se organizou os dados obtidos durante as etapas de observação e regência, sistematizando os resultados e reflexões em um documento final.

**VI – Apresentação dos resultados (5 horas):**
Por fim, preparou-se e apresentou-se um Banner com o objetivo de compartilhar as experiências vivenciadas no estágio, bem como os resultados alcançados e os desafios enfrentados, promovendo a troca de experiências entre os colegas de curso, professores orientadores e as demais pessoas presente no local.

Dessa forma, as atividades que compõem o {estagio_nome} permitiram aos acadêmicos compreender a realidade escolar, identificar dificuldades e potencialidades do ambiente educativo e aplicar, de forma concreta, os conteúdos teóricos estudados ao longo da graduação. Tal experiência foi de grande relevância para a formação docente, possibilitando um olhar crítico, reflexivo e propositivo sobre a prática pedagógica e os caminhos da educação contemporânea."""

            print(f"\n--- TEXTO GERADO ---\n{texto_final}\n--------------------")
            decisao_met = input("\n[S] Salvar | [P] Gerar nova versão | [N] Descartar e voltar: ").lower()
            if decisao_met == 's':
                dados['metodologia'] = texto_final
                salvar_dados(dados)
                print("\n✓ PROCEDIMENTOS METODOLÓGICOS salvos com sucesso!")
            elif decisao_met == 'p':
                print("\nGerando nova versão...\n")
                # Reinicia o loop de coleta de dados
                continue
            else:
                print("\nDescartado. Voltando ao menu.")

        else:
            while True:
                rel = input(f"Descreva o que deseja para {chave.upper()}:\n> ")
                if input("Confirma? (s/n): ").lower() == 's':
                    print(f"\nProcessando {chave.upper()}... Aguarde.")
                    res = chamar_minha_ia(f"Escreva {chave}: {rel}", modelo_ativo)
                    texto_l, refs_l = processar_e_extrair_refs(res, dados)
                    print(f"\n--- TEXTO GERADO ---\n{texto_l}\n--------------------")
                    decisao = input("\n[S] Salvar | [P] Gerar nova versão | [N] Descartar e voltar: ").lower()
                    if decisao == 's':
                        dados[chave] = (dados[chave] + "\n" + texto_l) if dados[chave] else texto_l
                        refs_atuais = dados['referencias'].split('\n')
                        for r in refs_l:
                            if r not in refs_atuais: refs_atuais.append(r)
                        dados['referencias'] = "\n".join(filter(None, sorted(refs_atuais)))
                        salvar_dados(dados)
                        print(f"\n✓ {chave.upper()} salvo com sucesso!")
                        break
                    elif decisao == 'p':
                        print("\nGerando nova versão...\n")
                        continue
                    else:
                        print("\nDescartado. Voltando ao menu.")
                        break
        salvar_dados(dados)

    elif op.lower() == "c":
        dados = configurar_capa(dados)
    elif op == "7": gerar_word_completo(dados)
    elif op == "8": break
    elif op == "9":
        if input("Deseja APAGAR TUDO para recomeçar? (s/n): ").lower() == 's':
            n_doc = f"Relatorio_{dados['autor'].replace(' ', '_') if dados['autor'] else 'Estagio'}.docx"
            if os.path.exists(n_doc): os.remove(n_doc)
            dados = {"titulo_trabalho":dados['titulo_trabalho'],"autor":dados['autor'],"orientador":dados['orientador'],"cidade":dados['cidade'],"ano":dados['ano'],"resumo":"","introducao":"","referencial":"","metodologia":"","resultados":"","conclusao":"","referencias":"","titulo_secao_2":"2. REFERENCIAL TEÓRICO"}
            salvar_dados(dados)
            print("\n✓ Reset concluído!")
