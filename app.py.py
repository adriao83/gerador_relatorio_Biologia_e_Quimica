import streamlit as st
import requests
import json
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="Assistente de Estágio UFAM",
    page_icon="🎓",
    layout="wide"
)

# --- CHAVE DE API ---
minha_chave = os.environ.get("GEMINI_API_KEY", "")

# --- ESTILO CSS ---
st.markdown("""
<style>
    .titulo-secao { font-size: 1.2rem; font-weight: bold; color: #1a5276; margin-bottom: 0.5rem; }
    .status-ok { color: #27ae60; font-weight: bold; }
    .status-vazio { color: #e74c3c; font-weight: bold; }
    .caixa-texto { background: #f8f9fa; border-left: 4px solid #1a5276; padding: 1rem; border-radius: 4px; margin: 0.5rem 0; font-size: 0.9rem; white-space: pre-wrap; }
    .header-ufam { background: #1a5276; color: white; padding: 1rem 2rem; border-radius: 8px; margin-bottom: 1.5rem; }
</style>
""", unsafe_allow_html=True)

# --- MOTOR DE IA ---
@st.cache_resource
def encontrar_melhor_modelo():
    if not minha_chave:
        return None
    url = f"https://generativelanguage.googleapis.com/v1beta/models?key={minha_chave}"
    try:
        res = requests.get(url).json()
        modelos = [m['name'] for m in res.get('models', []) if 'generateContent' in m.get('supportedGenerationMethods', [])]
        for pref in ['2.0-flash', '1.5-flash', 'flash', 'pro']:
            for m in modelos:
                if pref in m:
                    return m
        return modelos[0] if modelos else None
    except:
        return None

def chamar_ia(prompt, modelo):
    if not minha_chave or not modelo:
        return "Erro: Chave de API não configurada."
    url = f"https://generativelanguage.googleapis.com/v1beta/{modelo}:generateContent?key={minha_chave}"
    prompt_final = (f"{prompt} REGRAS: Inicie o texto diretamente. Use linguagem acadêmica e impessoal. "
                    f"IMPORTANTE: Se usar citações, ao final do texto, escreva a tag [REFS] e coloque as referências completas em ABNT.")
    corpo = {"contents": [{"parts": [{"text": prompt_final}]}]}
    try:
        res = requests.post(url, headers={'Content-Type': 'application/json'}, data=json.dumps(corpo))
        return res.json()['candidates'][0]['content']['parts'][0]['text']
    except:
        return "Erro na conexão com a IA."

def processar_refs(resposta):
    meses = ["jan.", "fev.", "mar.", "abr.", "mai.", "jun.", "jul.", "ago.", "set.", "out.", "nov.", "dez."]
    hoje = datetime.now()
    data_fmt = f"{hoje.day:02d} {meses[hoje.month-1]} {hoje.year}"
    texto_limpo = resposta
    refs = []
    if "[REFS]" in resposta:
        partes = resposta.split("[REFS]")
        texto_limpo = partes[0].strip()
        for r in partes[1].strip().split('\n'):
            if r.strip():
                r = r.replace('*', '').strip()
                if "Acesso em:" in r:
                    r = r.split("Acesso em:")[0] + f"Acesso em: {data_fmt}."
                refs.append(r)
    return texto_limpo, refs

# --- MEMÓRIA DA SESSÃO ---
def init_state():
    if 'dados' not in st.session_state:
        st.session_state.dados = {
            "titulo_trabalho": "", "autor": "", "orientador": "",
            "cidade": "", "ano": "", "resumo": "", "introducao": "",
            "referencial": "", "metodologia": "", "resultados": "",
            "conclusao": "", "referencias": "", "titulo_secao_2": "REFERENCIAL TEÓRICO"
        }
    if 'modelo' not in st.session_state:
        st.session_state.modelo = encontrar_melhor_modelo()
    if 'texto_gerado' not in st.session_state:
        st.session_state.texto_gerado = {}

init_state()
d = st.session_state.dados

# --- FUNÇÕES WORD ---
def add_page_number(paragraph):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    for tag, val in [('begin', None), ('instrText', 'PAGE'), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve')
            el.text = val
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        run._r.append(el)

def add_estilo(doc, texto, negrito=False, tam=12, alinhar=WD_ALIGN_PARAGRAPH.JUSTIFY,
               recuo_esq=0, recuo_primeira_linha=0, eh_simples=False):
    for trecho in texto.split('\n'):
        trecho = trecho.strip()
        if not trecho:
            continue

        if trecho in ["RESUMO", "SUMÁRIO"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(trecho)
            run.bold = True
            run.font.size, run.font.name = Pt(12), 'Times New Roman'
            continue

        eh_legenda = trecho.startswith("Figura 01:") or trecho.startswith("Fonte:") or trecho == "(LOCAL DA FOTO)"
        eh_etapa_romana = bool(re.match(r'^\*\*(I{1,3}V?|VI?I?I?|IV|IX|X)\s*[\u2013\-]', trecho))
        trecho_limpo = trecho.replace('**', '').strip()

        partes = trecho.split()
        primeira = partes[0] if partes else ""
        eh_num = primeira and primeira[0].isdigit() and "." in primeira
        nivel = None
        if eh_num and not eh_legenda:
            pontos = primeira.count('.')
            if pontos == 1:
                nivel = 1 if primeira.endswith('.') else 2
            elif pontos >= 2:
                nivel = 3

        p = doc.add_heading("", level=nivel) if nivel else doc.add_paragraph()

        if eh_legenda:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.first_line_indent = 0
        elif nivel:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = 0
            p.paragraph_format.space_before = Pt(24)
        elif eh_etapa_romana:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = 0
            p.paragraph_format.left_indent = 0
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(12)
        else:
            p.alignment = alinhar
            if recuo_esq > 0:
                p.paragraph_format.left_indent = Cm(recuo_esq)
            if recuo_primeira_linha > 0 and not eh_simples:
                p.paragraph_format.first_line_indent = Cm(recuo_primeira_linha)
            p.paragraph_format.line_spacing = 1.0 if eh_simples else 1.5

        texto_render = trecho_limpo if eh_etapa_romana else trecho
        run = p.add_run(texto_render)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)

        if eh_legenda:
            run.font.size = Pt(10)
            run.bold = trecho.startswith("Figura")
        elif eh_etapa_romana:
            run.font.size = Pt(tam)
            run.bold = True
        else:
            run.font.size = Pt(tam)
            if nivel == 1:
                run.bold = True
            elif nivel == 2:
                run.bold = run.italic = False
            else:
                run.bold = negrito

def gerar_word(dados):
    doc = Document()
    for s in doc.sections:
        s.top_margin, s.bottom_margin = Cm(3), Cm(2)
        s.left_margin, s.right_margin = Cm(3), Cm(2)

    # Capa
    add_estilo(doc, "UNIVERSIDADE FEDERAL DO AMAZONAS - UFAM\nINSTITUTO NATUREZA E CULTURA – INC\nCURSO DE LICENCIATURA EM CIÊNCIAS: BIOLOGIA E QUÍMICA",
               True, alinhar=WD_ALIGN_PARAGRAPH.CENTER, eh_simples=True)
    for _ in range(12): doc.add_paragraph()
    add_estilo(doc, dados['titulo_trabalho'].upper(), True, alinhar=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(16): doc.add_paragraph()
    add_estilo(doc, f"{dados['cidade'].upper()}\n{dados['ano']}", False, alinhar=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Folha de rosto
    add_estilo(doc, dados['autor'].upper(), True, alinhar=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(10): doc.add_paragraph()
    add_estilo(doc, dados['titulo_trabalho'].upper(), True, alinhar=WD_ALIGN_PARAGRAPH.CENTER)
    add_estilo(doc, "Relatório apresentado ao Instituto de Natureza e Cultura – INC/UFAM, como requisito parcial para obtenção de nota na disciplina de Estágio.",
               tam=10, recuo_esq=8, eh_simples=True)
    for _ in range(3): doc.add_paragraph()
    add_estilo(doc, f"Orientador (a): {dados['orientador']}", False, alinhar=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(8): doc.add_paragraph()
    add_estilo(doc, f"{dados['cidade'].upper()}\n{dados['ano']}", False, alinhar=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Resumo
    if dados.get('resumo'):
        add_estilo(doc, "RESUMO")
        add_estilo(doc, dados['resumo'], recuo_primeira_linha=1.25)
        doc.add_page_break()

    # Sumário
    add_estilo(doc, "SUMÁRIO")
    p_sum = doc.add_paragraph()
    p_sum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_sum = p_sum.add_run()
    fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin'); run_sum._r.append(fld)
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'), 'preserve'); ins.text = 'TOC \\o "1-3" \\h \\z \\u'; run_sum._r.append(ins)
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end'); run_sum._r.append(fld2)
    doc.add_page_break()

    # Seções
    new_sec = doc.add_section()
    new_sec.footer.is_linked_to_previous = False
    tit2 = dados.get('titulo_secao_2', 'REFERENCIAL TEÓRICO')
    if not tit2.startswith("2. "): tit2 = f"2. {tit2}"

    secoes = [
        ("1. INTRODUÇÃO", dados['introducao'], False),
        (tit2, dados['referencial'], False),
        ("3. PROCEDIMENTOS METODOLÓGICOS", dados['metodologia'], False),
        ("4. RESULTADOS E DISCUSSÕES", dados['resultados'], False),
        ("5. CONSIDERAÇÕES FINAIS", dados['conclusao'], False),
        ("6. REFERÊNCIAS BIBLIOGRÁFICAS", dados['referencias'], True),
    ]
    for titulo, conteudo, simples in secoes:
        if conteudo:
            add_estilo(doc, titulo, True)
            add_estilo(doc, conteudo, recuo_primeira_linha=1.25, eh_simples=simples)
            add_page_number(new_sec.footer.paragraphs[0])
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- INTERFACE ---
st.markdown('<div class="header-ufam"><h2>🎓 Assistente de Relatório de Estágio — UFAM/INC</h2><p>Licenciatura em Ciências: Biologia e Química</p></div>', unsafe_allow_html=True)

# Aviso se sem chave
if not minha_chave:
    st.error("⚠️ Chave de API não configurada! Vá em Settings > Secrets no Streamlit Cloud e adicione GEMINI_API_KEY.")

# --- ABAS ---
aba_capa, aba_resumo, aba_intro, aba_ref, aba_met, aba_result, aba_consid, aba_word = st.tabs([
    "📋 Capa", "📝 Resumo", "1️⃣ Introdução", "2️⃣ Referencial",
    "3️⃣ Metodologia", "4️⃣ Resultados", "5️⃣ Considerações", "📄 Gerar Word"
])

# ===================== ABA CAPA =====================
with aba_capa:
    st.subheader("📋 Capa e Folha de Rosto")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**CAPA**")
        d['titulo_trabalho'] = st.text_input("Título do Relatório", value=d['titulo_trabalho'])
        d['cidade'] = st.text_input("Cidade", value=d['cidade'])
        d['ano'] = st.text_input("Ano", value=d['ano'])
    with col2:
        st.markdown("**FOLHA DE ROSTO**")
        d['autor'] = st.text_input("Nome completo do Autor", value=d['autor'])
        d['orientador'] = st.text_input("Nome completo do Orientador(a)", value=d['orientador'])

    if st.button("💾 Salvar Capa e Folha de Rosto", type="primary"):
        st.success("✓ Capa e Folha de Rosto salvas!")

# ===================== FUNÇÃO GENÉRICA DE SEÇÃO =====================
def aba_secao(chave, titulo_secao, prompt_base, texto_ajuda):
    st.subheader(titulo_secao)
    conteudo_atual = d.get(chave, "")
    if conteudo_atual:
        st.markdown("**Conteúdo atual:**")
        st.markdown(f'<div class="caixa-texto">{conteudo_atual[:800]}{"..." if len(conteudo_atual) > 800 else ""}</div>', unsafe_allow_html=True)

    st.markdown("---")
    instrucao = st.text_area(texto_ajuda, height=120, key=f"instr_{chave}")

    col1, col2 = st.columns([1, 1])
    with col1:
        gerar = st.button(f"🤖 Gerar com IA", key=f"gerar_{chave}", type="primary")
    with col2:
        if conteudo_atual:
            limpar = st.button(f"🗑️ Limpar seção", key=f"limpar_{chave}")
            if limpar:
                d[chave] = ""
                st.rerun()

    if gerar:
        if not instrucao.strip():
            st.warning("Descreva o que deseja antes de gerar.")
        else:
            with st.spinner("Gerando texto com IA... Aguarde."):
                prompt = f"{prompt_base}: {instrucao}"
                res = chamar_ia(prompt, st.session_state.modelo)
                texto_l, refs_l = processar_refs(res)
                st.session_state.texto_gerado[chave] = (texto_l, refs_l)

    if chave in st.session_state.texto_gerado:
        texto_l, refs_l = st.session_state.texto_gerado[chave]
        st.markdown("**Texto gerado:**")
        st.markdown(f'<div class="caixa-texto">{texto_l}</div>', unsafe_allow_html=True)
        if refs_l:
            st.markdown(f"**Referências detectadas ({len(refs_l)}):** serão adicionadas automaticamente.")

        col_s, col_p, col_n = st.columns(3)
        with col_s:
            if st.button("✅ Salvar", key=f"salvar_{chave}", type="primary"):
                d[chave] = (d[chave] + "\n" + texto_l).strip() if d[chave] else texto_l
                refs_atuais = [r for r in d['referencias'].split('\n') if r.strip()]
                for r in refs_l:
                    if r not in refs_atuais:
                        refs_atuais.append(r)
                d['referencias'] = "\n".join(sorted(filter(None, refs_atuais)))
                del st.session_state.texto_gerado[chave]
                st.success(f"✓ {titulo_secao} salvo com sucesso!")
                st.rerun()
        with col_p:
            if st.button("🔄 Nova versão", key=f"nova_{chave}"):
                del st.session_state.texto_gerado[chave]
                st.rerun()
        with col_n:
            if st.button("❌ Descartar", key=f"descartar_{chave}"):
                del st.session_state.texto_gerado[chave]
                st.rerun()

# ===================== ABAS DE SEÇÕES =====================
with aba_resumo:
    aba_secao("resumo", "📝 Resumo", "Escreva o resumo acadêmico do relatório de estágio",
              "Descreva o que deseja no resumo (tema, objetivos, metodologia, resultados):")

with aba_intro:
    aba_secao("introducao", "1️⃣ Introdução", "Escreva a introdução do relatório de estágio",
              "Descreva o contexto do seu estágio para a IA desenvolver a introdução:")

with aba_ref:
    st.subheader("2️⃣ Referencial Teórico")
    tit2 = st.text_input("Título principal da Seção 2:", value=d.get('titulo_secao_2', 'REFERENCIAL TEÓRICO'))
    d['titulo_secao_2'] = tit2

    conteudo_ref = d.get('referencial', '')
    if conteudo_ref:
        st.markdown("**Conteúdo atual:**")
        st.markdown(f'<div class="caixa-texto">{conteudo_ref[:600]}{"..." if len(conteudo_ref) > 600 else ""}</div>', unsafe_allow_html=True)
        if st.button("🗑️ Limpar Referencial"):
            d['referencial'] = ''
            st.rerun()

    st.markdown("---")
    for i in ["2.1", "2.2", "2.3"]:
        with st.expander(f"📌 Subseção {i}", expanded=(i == "2.1")):
            sub_titulo = st.text_input(f"Título do subtítulo {i}:", key=f"tit_{i}")
            sub_instrucao = st.text_area(f"O que a IA escreve para '{sub_titulo or i}'?", height=100, key=f"instr_{i}")

            if st.button(f"🤖 Gerar {i}", key=f"gerar_{i}", type="primary"):
                if not sub_titulo.strip():
                    st.warning("Digite o título do subtítulo antes de gerar.")
                elif not sub_instrucao.strip():
                    st.warning("Descreva o conteúdo desejado.")
                else:
                    with st.spinner(f"Gerando {i}... Aguarde."):
                        prompt = (
                            f"Escreva um texto acadêmico completo com no mínimo 4 parágrafos bem desenvolvidos "
                            f"para a subseção '{i} {sub_titulo}', que faz parte do Referencial Teórico sobre o tema geral "
                            f"'{tit2}'. Instruções adicionais: {sub_instrucao}. "
                            f"NÃO escreva o título no início. Escreva APENAS o texto corrido em linguagem acadêmica, "
                            f"com citações de autores e embasamento científico."
                        )
                        res = chamar_ia(prompt, st.session_state.modelo)
                        texto_l, refs_l = processar_refs(res)
                        linhas = texto_l.split('\n')
                        if linhas[0].strip().lower().startswith(i) and len(linhas[0]) < 100:
                            texto_l = '\n'.join(linhas[1:]).strip()
                        st.session_state.texto_gerado[f"ref_{i}"] = (sub_titulo, texto_l, refs_l)

            if f"ref_{i}" in st.session_state.texto_gerado:
                sub_t, texto_l, refs_l = st.session_state.texto_gerado[f"ref_{i}"]
                st.markdown("**Texto gerado:**")
                st.markdown(f'<div class="caixa-texto">{texto_l}</div>', unsafe_allow_html=True)
                if refs_l:
                    st.markdown(f"**{len(refs_l)} referência(s) detectada(s)**")

                col_s, col_p, col_n = st.columns(3)
                with col_s:
                    if st.button(f"✅ Salvar {i}", key=f"salvar_ref_{i}", type="primary"):
                        d['referencial'] = (d['referencial'] + f"\n{i} {sub_t.upper()}\n{texto_l}\n").strip()
                        refs_atuais = [r for r in d['referencias'].split('\n') if r.strip()]
                        for r in refs_l:
                            if r not in refs_atuais: refs_atuais.append(r)
                        d['referencias'] = "\n".join(sorted(filter(None, refs_atuais)))
                        del st.session_state.texto_gerado[f"ref_{i}"]
                        st.success(f"✓ Subseção {i} salva!")
                        st.rerun()
                with col_p:
                    if st.button(f"🔄 Nova versão {i}", key=f"nova_ref_{i}"):
                        del st.session_state.texto_gerado[f"ref_{i}"]
                        st.rerun()
                with col_n:
                    if st.button(f"❌ Descartar {i}", key=f"desc_ref_{i}"):
                        del st.session_state.texto_gerado[f"ref_{i}"]
                        st.rerun()

with aba_met:
    st.subheader("3️⃣ Procedimentos Metodológicos")
    col1, col2 = st.columns(2)
    with col1:
        estagio_nome = st.text_input("Nome do Estágio:", value="Estágio Supervisionado no Ensino de Química")
        disciplina = st.text_input("Disciplina:", value="Química")
        escola = st.text_input("Nome da Escola:", value="Escola Estadual Indígena Professor Gildo Sampaio Megatanücü")
        local_escola = st.text_input("Localização:", value="na comunidade Aldeia Indígena Filadelfia, Benjamin Constant-AM")
    with col2:
        legenda_fig = st.text_input("Legenda da Figura 01:", value="Parte da frente da Escola")
        fonte_fig = st.text_input("Fonte da Figura:", value="Autor, 2025")
        turma_int = st.text_input("Turma da Intervenção:", value="101")
        tema_int = st.text_input("Tema da Intervenção:", value="ODS 15 Vida Terrestre")

    relato_regencia = st.text_area("Descreva a dinâmica/jogo da intervenção:", height=100)

    if st.button("🤖 Gerar Procedimentos Metodológicos", type="primary"):
        with st.spinner("Gerando texto com IA..."):
            prompt_reg = f"Com base no relato: {relato_regencia}. Escreva um parágrafo acadêmico curto para a seção de regência, sem títulos."
            texto_reg = chamar_ia(prompt_reg, st.session_state.modelo)

            texto_met = f"""O "{estagio_nome}" compreende atividades de observação no contexto escolar, bem como o acompanhamento e a prática da docência na disciplina de "{disciplina}" do 1º ao 3º ano do Ensino Médio. Inclui ainda o desenvolvimento do Planejamento de Ensino, Avaliação, Regência e Intervenção.

Na Figura 01 podemos observar a "{escola}".
Figura 01: {legenda_fig}
(LOCAL DA FOTO)
Fonte: {fonte_fig}.

A carga horária total do estágio é de 135 horas, distribuídas em seis (6) etapas, conforme descrito a seguir:

**I – Aulas Teóricas (45 horas):**
Inicialmente, foram ministradas aulas introdutórias abordando as normativas que regem o estágio no curso de Licenciatura em Ciências: Biologia e Química. Ao longo dos encontros teóricos, discutiram-se temas essenciais à formação docente, tais como: a relevância do estágio na formação de professores, a articulação entre teoria e prática, e os fundamentos da formação inicial de educadores. Durante essa fase, realizaram-se o preenchimento dos documentos principais do estágio, além de métodos pedagógicos para utilizar na prática, com o propósito de fixar os conteúdos explorados em sala.

**II – Observação Escolar (30 horas):**
Esta etapa teve como objetivo a inserção dos licenciandos na realidade escolar. Consistiu na análise do espaço físico, da estrutura organizacional, do funcionamento da instituição, da proposta pedagógica e das relações interpessoais estabelecidas no ambiente educacional. A escola escolhida foi a "{escola}", localizada {local_escola}. A observação foi realizada nos seguintes espaços: setores administrativos, biblioteca, laboratório de Ciências, refeitório, áreas externas e salas de aula, com dedicação de 5 horas para cada setor.

**III – Planejamento Pedagógico (20 horas):**
A etapa de planejamento consistiu na elaboração de atividades que seriam aplicadas durante a regência. Após a definição das turmas em que os estagiários atuariam, o professor responsável pela disciplina de "{disciplina}" indicou os conteúdos a serem trabalhados. Com base nisso, foram construídos o plano de estágio, os planos de aula e o projeto de intervenção pedagógica.

**IV – Regência e Intervenção (20 horas):**
Neste momento, realizaram-se 15 horas de regência efetiva em sala e 5 horas destinadas à execução do plano de intervenção. A atuação docente ocorreu com todas as turmas do 101 ao 302 do Ensino Médio. As aulas foram conduzidas por meio de metodologias diversificadas, incluindo o uso do quadro branco, vídeos didáticos e interação com os estudantes. A intervenção foi aplicada na turma {turma_int}, com o tema {tema_int}. {texto_reg}

**V – Elaboração do Relatório Final (15 horas):**
Essa fase consistiu na construção do relatório descritivo do estágio. Foram realizadas pesquisas bibliográficas sobre a temática "{d['titulo_trabalho'] if d['titulo_trabalho'] else 'formação docente'}", as quais fundamentaram teoricamente a escrita do relatório. Também se organizou os dados obtidos durante as etapas de observação e regência, sistematizando os resultados e reflexões em um documento final.

**VI – Apresentação dos resultados (5 horas):**
Por fim, preparou-se e apresentou-se um Banner com o objetivo de compartilhar as experiências vivenciadas no estágio, bem como os resultados alcançados e os desafios enfrentados, promovendo a troca de experiências entre os colegas de curso e professores orientadores.

Dessa forma, as atividades que compõem o {estagio_nome} permitiram aos acadêmicos compreender a realidade escolar, identificar dificuldades e potencialidades do ambiente educativo e aplicar, de forma concreta, os conteúdos teóricos estudados ao longo da graduação."""

            st.session_state.texto_gerado['metodologia'] = (texto_met, [])

    if 'metodologia' in st.session_state.texto_gerado:
        texto_l, _ = st.session_state.texto_gerado['metodologia']
        st.markdown("**Texto gerado:**")
        st.markdown(f'<div class="caixa-texto">{texto_l}</div>', unsafe_allow_html=True)

        col_s, col_p, col_n = st.columns(3)
        with col_s:
            if st.button("✅ Salvar Metodologia", type="primary"):
                d['metodologia'] = texto_l
                del st.session_state.texto_gerado['metodologia']
                st.success("✓ Procedimentos Metodológicos salvos!")
                st.rerun()
        with col_p:
            if st.button("🔄 Nova versão"):
                del st.session_state.texto_gerado['metodologia']
                st.rerun()
        with col_n:
            if st.button("❌ Descartar"):
                del st.session_state.texto_gerado['metodologia']
                st.rerun()

with aba_result:
    aba_secao("resultados", "4️⃣ Resultados e Discussões",
              "Escreva a seção de Resultados e Discussões do relatório de estágio",
              "Descreva os resultados e experiências vivenciadas no estágio:")

with aba_consid:
    aba_secao("conclusao", "5️⃣ Considerações Finais",
              "Escreva as Considerações Finais do relatório de estágio",
              "Descreva suas reflexões finais sobre o estágio:")

# ===================== ABA GERAR WORD =====================
with aba_word:
    st.subheader("📄 Gerar Relatório Word")

    campos = {
        "Capa / Folha de Rosto": all(d.get(c) for c in ['titulo_trabalho', 'autor', 'orientador', 'cidade', 'ano']),
        "Resumo": bool(d.get('resumo')),
        "Introdução": bool(d.get('introducao')),
        "Referencial Teórico": bool(d.get('referencial')),
        "Procedimentos Metodológicos": bool(d.get('metodologia')),
        "Resultados e Discussões": bool(d.get('resultados')),
        "Considerações Finais": bool(d.get('conclusao')),
        "Referências Bibliográficas": bool(d.get('referencias')),
    }

    col1, col2 = st.columns(2)
    for i, (nome, ok) in enumerate(campos.items()):
        col = col1 if i % 2 == 0 else col2
        with col:
            status = "✅" if ok else "⬜"
            st.markdown(f"{status} {nome}")

    st.markdown("---")
    if not campos["Capa / Folha de Rosto"]:
        st.warning("⚠️ Preencha a Capa e Folha de Rosto antes de gerar o Word.")
    else:
        if st.button("📥 Gerar e Baixar Word", type="primary"):
            with st.spinner("Gerando documento Word..."):
                try:
                    buf = gerar_word(d)
                    nome_arquivo = f"Relatorio_{d['autor'].replace(' ', '_') if d['autor'] else 'Estagio'}.docx"
                    st.download_button(
                        label="⬇️ Clique aqui para baixar o Word",
                        data=buf,
                        file_name=nome_arquivo,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success("✓ Word gerado com sucesso!")
                except Exception as e:
                    st.error(f"Erro ao gerar Word: {e}")
